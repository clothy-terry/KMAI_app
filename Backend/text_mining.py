import os
import fitz  # PyMuPDF
import docx
import openpyxl
import json
import re
import shutil
from docx import Document
from spire.doc import *
from spire.doc.common import *

def process_documents(user_doc_dir = "user_documents", output_file = "user_documents/user_doc_text.jsonl"):

    def clean_text(text):
        # Lowercase the text
        text = text.lower()
        
        # Remove URLs
        text = re.sub(r"http\S+|www\S+|https\S+", '', text, flags=re.MULTILINE)
        
        # Remove special characters and extra whitespace
        text = ''.join(i if ord(i) < 128 else ' ' for i in text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text

    import pytesseract
    from PIL import Image

    # Define the path to the tesseract executable
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Update this path

    """
    from pdf2image import convert_from_path
    def extract_text_from_scanned_pdf2(pdf_path):
        pages = convert_from_path(pdf_path, 500)
        for pageNum,imgBlob in enumerate(pages):
            text = pytesseract.image_to_string(imgBlob,lang='eng')
        return text
    """

    def extract_text_from_scanned_pdf(pdf_path):
        document = fitz.open(pdf_path)
        text = ""

        # Iterate over each page
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            pix = page.get_pixmap()  # Get the page as a pixmap

            # Convert pixmap to an image
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Perform OCR on the image
            text += pytesseract.image_to_string(image,lang='eng')

        return text


    def extract_text_from_pdf(pdf_path):
        text = ""
        document = fitz.open(pdf_path)
        for page in document:
            text += page.get_text()
        return text

    def extract_text_from_docx(docx_path):
        text = ""
        document = docx.Document(docx_path)
        for para in document.paragraphs:
            text += para.text + "\n"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text


    def extract_text_from_doc(file_path):
            # Create a Document object
            document = Document()
            # Load a Word document
            document.LoadFromFile(file_path)

            # Extract the text of the document
            document_text = document.GetText()

            document.Close()
            
            return document_text


    def extract_text_from_xlsx(xlsx_path):
        text = ""
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for row in worksheet.iter_rows(values_only=True):
                row_text = " ".join([str(cell) for cell in row if cell is not None])
                text += row_text + "\n"
        return text

    def collect_text_from_files(directory):
        
        failed_dir = "..\Skipped_files"
        os.makedirs(failed_dir, exist_ok=True)

        skipped_files = []
        collected_text = []
        file_id = 1
        total_files = sum([len(files) for r, d, files in os.walk(directory)])
        print(f"Total files to process: {total_files}")
        for root, _, files in os.walk(directory):
            for file in files:
                is_skipped = False
                file_path = os.path.join(root, file)
                folder_name = os.path.basename(root)
                if file.lower().endswith(".pdf"):
                    text = extract_text_from_pdf(file_path)
                    if not text:
                        text = extract_text_from_scanned_pdf(file_path)
                elif file.lower().endswith(".docx"):
                    text = extract_text_from_docx(file_path)
                elif file.lower().endswith(".doc"):
                    text = extract_text_from_doc(file_path)
                elif file.lower().endswith(".xlsx"):
                    text = extract_text_from_xlsx(file_path)
                else:
                    skipped_files.append(file_path)
                    is_skipped = True
                if not text and not is_skipped: # fail to process required documents
                    print("failt to collect: {file_path}")
                    return
                    skipped_files.append(file_path)
                cleaned_text = clean_text(text)
                collected_text.append({
                    "id": f"{folder_name} / {file}",
                    "text": cleaned_text
                })
                print(f"Processed file {file_id} of {total_files}: {folder_name} / {file}")
                file_id += 1
        for file_path in skipped_files:
            shutil.copy2(file_path, failed_dir)
        return collected_text, skipped_files

    def save_text_to_jsonl(collected_text, output_file):
        with open(output_file, 'w', encoding='utf-8') as f:
            for entry in collected_text:
                json.dump(entry, f, ensure_ascii=False)
                f.write('\n')


    collected_text, skipped_files = collect_text_from_files(user_doc_dir)
    print('skipped files', skipped_files)
    save_text_to_jsonl(collected_text, output_file)

    return collected_text, skipped_files

